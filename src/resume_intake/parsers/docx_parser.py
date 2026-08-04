from __future__ import annotations

from pathlib import Path

from docx import Document

from resume_intake.parsers.base import ResumeParser
from resume_intake.types import ParsedDocument


class DocxResumeParser(ResumeParser):
    supported_extensions = (".docx",)
    parser_name = "python-docx"

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            document = Document(str(file_path))
        except Exception as exc:
            raise ValueError(f"Unable to read DOCX file: {file_path}") from exc

        chunks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        return ParsedDocument(text="\n".join(chunks).strip(), metadata={})
